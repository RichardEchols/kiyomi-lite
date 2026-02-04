"""
Kiyomi Lite — Tool Layer

Defines tools the AI can call silently:
- web_search (DuckDuckGo HTML)
- read_url (via url_reader.fetch_url)
- read_file (restricted to ~/.kiyomi/)
- run_code (python3 subprocess w/ timeout)
- remember (append learned fact to memory)
- create_file (create .docx or .txt documents)
- send_email (send email via Gmail SMTP)
- analyze_image (analyze an image with AI vision)
"""

from __future__ import annotations

import base64
import json
import smtplib
import subprocess
import urllib.parse
import urllib.request
from dataclasses import dataclass
from datetime import datetime
from email.mime.text import MIMEText
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from engine.config import CONFIG_DIR, MEMORY_DIR, ensure_dirs, load_config
from url_reader import fetch_url

FILES_DIR = CONFIG_DIR / "files"


TOOLS = [
    {
        "name": "web_search",
        "description": (
            "Search the web for current information. Use when user asks about news, facts, "
            "current events, or anything you don't know."
        ),
        "parameters": {
            "query": {"type": "string", "description": "Search query"},
        },
    },
    {
        "name": "read_url",
        "description": (
            "Read and extract text content from a web URL. Use when user sends a link or "
            "you need to read a specific webpage."
        ),
        "parameters": {
            "url": {"type": "string", "description": "The URL to read"},
        },
    },
    {
        "name": "read_file",
        "description": (
            "Read the contents of a text file. Use when user mentions a file or you need "
            "to read stored data."
        ),
        "parameters": {
            "path": {"type": "string", "description": "Path to the file"},
        },
    },
    {
        "name": "run_code",
        "description": (
            "Run a Python code snippet and return the output. Use for calculations, data "
            "processing, date math, unit conversions, or any computation."
        ),
        "parameters": {
            "code": {"type": "string", "description": "Python code to execute"},
        },
    },
    {
        "name": "remember",
        "description": (
            "Save an important fact about the user to long-term memory. Use when you learn "
            "something important about the user that should be remembered."
        ),
        "parameters": {
            "fact": {"type": "string", "description": "The fact to remember"},
            "category": {
                "type": "string",
                "description": "Category: identity, family, work, health, preferences, goals, schedule, other",
            },
        },
    },
    {
        "name": "create_file",
        "description": (
            "Create a document file (.docx or .txt). Use when user asks you to write a resume, "
            "letter, report, or any document they can download. Write content in Markdown "
            "(# headings, **bold**, *italic*, - bullets). Set format for document type."
        ),
        "parameters": {
            "filename": {"type": "string", "description": "Filename with extension (.docx or .txt)"},
            "content": {"type": "string", "description": "The text content to write into the file (use Markdown formatting)"},
            "title": {"type": "string", "description": "Optional document title for .docx files"},
            "format": {"type": "string", "description": "Document format: 'general', 'resume', 'letter', 'report'. Default: 'general'"},
        },
    },
    {
        "name": "send_email",
        "description": (
            "Send an email. Use when user asks you to email someone. "
            "Requires email to be configured first."
        ),
        "parameters": {
            "to": {"type": "string", "description": "Recipient email address"},
            "subject": {"type": "string", "description": "Email subject line"},
            "body": {"type": "string", "description": "Email body text"},
        },
    },
    {
        "name": "analyze_image",
        "description": (
            "Analyze an image file. Use when user sends a photo or asks about an image."
        ),
        "parameters": {
            "image_path": {"type": "string", "description": "Path to the image file"},
            "question": {"type": "string", "description": "What to analyze about the image"},
        },
    },
    {
        "name": "check_calendar",
        "description": (
            "Check today's calendar events and schedule. Use when user asks about their "
            "schedule, what's on their calendar, what they have today, or upcoming events."
        ),
        "parameters": {
            "days": {
                "type": "integer",
                "description": "Number of days to look ahead. 0 or 1 = today only (default). Use 7 for this week.",
            },
        },
    },
    {
        "name": "create_event",
        "description": (
            "Create a new calendar event. Use when user asks to schedule something, "
            "add an event, book time, or set an appointment."
        ),
        "parameters": {
            "title": {"type": "string", "description": "Event title (e.g., 'Dentist Appointment')"},
            "start": {
                "type": "string",
                "description": "Start time in ISO format: 'YYYY-MM-DDTHH:MM:SS' for timed events or 'YYYY-MM-DD' for all-day",
            },
            "end": {
                "type": "string",
                "description": "End time in ISO format: 'YYYY-MM-DDTHH:MM:SS' for timed events or 'YYYY-MM-DD' for all-day",
            },
            "description": {"type": "string", "description": "Optional event description or notes"},
            "location": {"type": "string", "description": "Optional event location"},
        },
    },
    {
        "name": "find_free_time",
        "description": (
            "Find open time slots on a given day. Use when user asks when they're free, "
            "wants to find time for something, or asks about availability."
        ),
        "parameters": {
            "date": {
                "type": "string",
                "description": "Date to check in YYYY-MM-DD format. Defaults to today if empty.",
            },
        },
    },
    {
        "name": "check_spending",
        "description": (
            "Check spending summary from the user's connected bank account. "
            "Use when user asks about spending, budget, expenses, how much they spent, "
            "or financial overview. Shows spending by category and top merchants."
        ),
        "parameters": {
            "days": {"type": "integer", "description": "Number of days to look back (default 30). Use 7 for this week, 30 for this month."},
        },
    },
    {
        "name": "check_balances",
        "description": (
            "Check current bank account balances and net worth. "
            "Use when user asks about their balance, how much money they have, "
            "or their net worth."
        ),
        "parameters": {},
    },
    {
        "name": "check_category_spending",
        "description": (
            "Check spending for a specific category like food, shopping, transportation, etc. "
            "Use when user asks 'how much did I spend on food?' or similar category questions."
        ),
        "parameters": {
            "category": {"type": "string", "description": "Spending category to check (e.g., Food and Drink, Shopping, Transportation, Entertainment)"},
            "days": {"type": "integer", "description": "Number of days to look back (default 30)"},
        },
    },
]


def _truncate(text: str, max_chars: int) -> str:
    if text is None:
        return ""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "..."


# ============================================================
# web_search (DuckDuckGo HTML)
# ============================================================


@dataclass(frozen=True)
class _SearchResult:
    title: str
    url: str
    snippet: str


class _DuckDuckGoHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.results: list[_SearchResult] = []
        self._in_title = False
        self._in_snippet = False
        self._current_title_parts: list[str] = []
        self._current_snippet_parts: list[str] = []
        self._current_url: str | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]):
        if tag != "a" and tag != "div":
            return

        attrs_dict = {k: v for k, v in attrs}
        class_attr = attrs_dict.get("class") or ""

        if tag == "a" and "result__a" in class_attr:
            self._flush_current()
            self._in_title = True
            self._current_url = (attrs_dict.get("href") or "").strip()
            return

        # DuckDuckGo uses either <a class="result__snippet"> or <div class="result__snippet">
        if "result__snippet" in class_attr:
            self._in_snippet = True

    def handle_endtag(self, tag: str):
        if tag == "a" and self._in_title:
            self._in_title = False
            return
        if self._in_snippet and tag in ("a", "div"):
            self._in_snippet = False

    def handle_data(self, data: str):
        if self._in_title:
            self._current_title_parts.append(data)
        elif self._in_snippet:
            self._current_snippet_parts.append(data)

    def close(self):
        self._flush_current()
        super().close()

    def _flush_current(self):
        title = " ".join(p.strip() for p in self._current_title_parts).strip()
        snippet = " ".join(p.strip() for p in self._current_snippet_parts).strip()
        url = (self._current_url or "").strip()

        if title and url:
            self.results.append(_SearchResult(title=title, url=url, snippet=snippet))

        self._current_title_parts = []
        self._current_snippet_parts = []
        self._current_url = None


def web_search(query: str) -> str:
    query = (query or "").strip()
    if not query:
        return "Error: missing search query."

    ddg_url = "https://html.duckduckgo.com/html/?" + urllib.parse.urlencode({"q": query})
    req = urllib.request.Request(
        ddg_url,
        headers={
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
            "Accept": "text/html,application/xhtml+xml,*/*",
        },
    )

    try:
        with urllib.request.urlopen(req, timeout=10) as resp:
            raw = resp.read(512_000)
            html = raw.decode("utf-8", errors="replace")
    except Exception as e:
        return f"Error: web_search failed: {type(e).__name__}: {str(e)[:200]}"

    parser = _DuckDuckGoHTMLParser()
    try:
        parser.feed(html)
        parser.close()
    except Exception as e:
        return f"Error: web_search parse failed: {type(e).__name__}: {str(e)[:200]}"

    results = parser.results[:5]
    if not results:
        return "No results found."

    lines: list[str] = []
    for idx, r in enumerate(results, start=1):
        lines.append(f"{idx}. {r.title}\n   {r.url}")
        if r.snippet:
            lines.append(f"   {r.snippet}")
    return "\n".join(lines).strip()


# ============================================================
# read_url
# ============================================================


def read_url(url: str) -> str:
    url = (url or "").strip()
    if not url:
        return "Error: missing url."

    content = fetch_url(url, max_chars=3000, timeout=10)
    if not content:
        return f"Error: could not read URL: {url}"
    return _truncate(content, 3000)


# ============================================================
# read_file (restricted to ~/.kiyomi/)
# ============================================================


_ALLOWED_EXTS = {".txt", ".md", ".csv", ".json"}


def _resolve_under_kiyomi(path: str) -> Path:
    expanded = Path(path).expanduser()
    resolved = expanded.resolve()
    allowed_root = CONFIG_DIR.resolve()
    if not resolved.is_relative_to(allowed_root):
        raise PermissionError("read_file only allows paths under ~/.kiyomi/")
    return resolved


def read_file(path: str) -> str:
    ensure_dirs()
    path = (path or "").strip()
    if not path:
        return "Error: missing path."

    try:
        file_path = _resolve_under_kiyomi(path)
    except Exception as e:
        return f"Error: {str(e)[:200]}"

    if file_path.suffix.lower() not in _ALLOWED_EXTS:
        return f"Error: unsupported file type '{file_path.suffix}'. Allowed: {', '.join(sorted(_ALLOWED_EXTS))}"

    if not file_path.exists():
        return f"Error: file not found: {file_path}"

    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read(64_000)
    except Exception as e:
        return f"Error: failed to read file: {type(e).__name__}: {str(e)[:200]}"

    return _truncate(content, 5000)


# ============================================================
# run_code
# ============================================================


def run_code(code: str) -> str:
    code = code or ""
    try:
        proc = subprocess.run(
            ["python3", "-c", code],
            capture_output=True,
            text=True,
            timeout=10,
        )
    except subprocess.TimeoutExpired:
        return "Error: code execution timed out after 10 seconds."
    except Exception as e:
        return f"Error: run_code failed: {type(e).__name__}: {str(e)[:200]}"

    stdout = (proc.stdout or "").strip()
    stderr = (proc.stderr or "").strip()

    out = ""
    if stdout:
        out += stdout
    if stderr:
        out += ("\n" if out else "") + stderr
    if not out:
        out = "No output."

    return _truncate(out, 2000)


# ============================================================
# remember
# ============================================================


_ALLOWED_CATEGORIES = {
    "identity",
    "family",
    "work",
    "health",
    "preferences",
    "goals",
    "schedule",
    "other",
}


def remember(fact: str, category: str) -> str:
    ensure_dirs()
    fact = (fact or "").strip()
    category = (category or "").strip().lower() or "other"

    if not fact:
        return "Error: missing fact."
    if category not in _ALLOWED_CATEGORIES:
        category = "other"

    facts_file = MEMORY_DIR / "learned_facts.md"
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    line = f"- [{ts}] [{category}] {fact}\n"

    try:
        if facts_file.exists():
            existing = facts_file.read_text(encoding="utf-8", errors="replace")
            if not existing.endswith("\n"):
                existing += "\n"
        else:
            existing = "# Things I've Learned\n\n"
        facts_file.write_text(existing + line, encoding="utf-8")
    except Exception as e:
        return f"Error: remember failed: {type(e).__name__}: {str(e)[:200]}"

    return "Saved to memory."


# ============================================================
# create_file (with Markdown → docx conversion)
# ============================================================


def _add_formatted_text(paragraph, text: str):
    """Add text to a paragraph with **bold** and *italic* inline formatting."""
    import re

    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(content: str, title: str = "", doc_format: str = "general"):
    """Convert Markdown text to a formatted Word document.

    Handles:
    - # Heading 1, ## Heading 2, ### Heading 3
    - **bold text**, *italic text*
    - - bullet points (unordered list)
    - 1. numbered items (ordered list)
    - --- horizontal rule / section break
    - Regular paragraphs

    Format presets:
    - 'resume': Narrow margins, compact spacing, name prominent
    - 'letter': Standard margins (1 inch)
    - 'report': Wider margins, generous spacing
    - 'general': Clean defaults
    """
    import re

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required for .docx creation. Install with: pip install python-docx")

    doc = DocxDocument()

    # --- Set margins based on format ---
    for section in doc.sections:
        if doc_format == "resume":
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        elif doc_format == "letter":
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        elif doc_format == "report":
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        else:  # general
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)

    # --- Optional title ---
    if title:
        heading = doc.add_heading(title, level=0)
        if doc_format == "resume":
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Parse Markdown line by line ---
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip blank lines
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            if doc_format == "resume":
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            if doc_format == "resume":
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Horizontal rule / section break
        elif stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.add_run("_" * 50)
        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)
        # Numbered list
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)
        # Regular paragraph — collect consecutive non-special lines
        else:
            para_lines = [stripped]
            while (
                i + 1 < len(lines)
                and lines[i + 1].strip()
                and not lines[i + 1].strip().startswith(("#", "---", "***", "___"))
                and not re.match(r"^[-*]\s", lines[i + 1].strip())
                and not re.match(r"^\d+\.\s", lines[i + 1].strip())
            ):
                i += 1
                para_lines.append(lines[i].strip())
            p = doc.add_paragraph()
            _add_formatted_text(p, " ".join(para_lines))

        i += 1

    return doc


def _strip_markdown(text: str) -> str:
    """Strip Markdown formatting for plain-text output."""
    import re

    # Remove heading markers
    text = re.sub(r"^#{1,3}\s+", "", text, flags=re.MULTILINE)
    # Remove bold markers
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    # Remove italic markers
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    # Remove horizontal rules
    text = re.sub(r"^(---|\*\*\*|___)\s*$", "", text, flags=re.MULTILINE)
    return text


def create_file(filename: str, content: str, title: str = "", format: str = "general") -> str:
    filename = (filename or "").strip()
    content = content or ""
    title = (title or "").strip()
    format = (format or "general").strip().lower()

    if not filename:
        return "Error: missing filename."
    if format not in ("general", "resume", "letter", "report"):
        format = "general"

    FILES_DIR.mkdir(parents=True, exist_ok=True)
    file_path = FILES_DIR / filename
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".docx":
            doc = markdown_to_docx(content, title, doc_format=format)
            doc.save(str(file_path))
        elif suffix == ".txt":
            # Strip markdown formatting for plain text
            plain = _strip_markdown(content)
            file_path.write_text(plain, encoding="utf-8")
        else:
            # Other extensions — write as-is
            file_path.write_text(content, encoding="utf-8")

        return str(file_path)
    except Exception as e:
        return f"Error: create_file failed: {type(e).__name__}: {str(e)[:200]}"


# ============================================================
# send_email
# ============================================================


def send_email(to: str, subject: str, body: str) -> str:
    to = (to or "").strip()
    subject = (subject or "").strip()
    body = body or ""

    if not to:
        return "Error: missing recipient email address."

    config = load_config()
    email_address = config.get("email_address", "")
    email_password = config.get("email_password", "")

    if not email_address or not email_password:
        return "Email not set up yet. Ask your user to configure email in settings."

    try:
        msg = MIMEText(body)
        msg["Subject"] = subject
        msg["From"] = email_address
        msg["To"] = to

        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(email_address, email_password)
            server.send_message(msg)

        return f"Email sent to {to}"
    except Exception as e:
        return f"Error: send_email failed: {type(e).__name__}: {str(e)[:200]}"


# ============================================================
# analyze_image
# ============================================================


def analyze_image(image_path: str, question: str = "") -> str:
    image_path = (image_path or "").strip()
    question = (question or "What's in this image? Describe it.").strip()

    if not image_path:
        return "Error: missing image_path."

    path = Path(image_path)
    if not path.exists():
        return f"Error: image file not found: {image_path}"

    try:
        image_bytes = path.read_bytes()
    except Exception as e:
        return f"Error: could not read image: {type(e).__name__}: {str(e)[:200]}"

    config = load_config()
    provider = config.get("provider", "gemini")

    try:
        if provider == "gemini":
            api_key = config.get("gemini_key", "")
            if not api_key:
                return "Error: Gemini API key not configured."
            import google.generativeai as genai

            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.0-flash")
            uploaded = genai.upload_file(str(path))
            response = model.generate_content([question, uploaded])
            return response.text or "No response from Gemini."

        elif provider == "anthropic":
            api_key = config.get("anthropic_key", "")
            if not api_key:
                return "Error: Anthropic API key not configured."
            import anthropic

            client = anthropic.Anthropic(api_key=api_key)
            b64 = base64.standard_b64encode(image_bytes).decode("utf-8")
            # Detect media type
            suffix = path.suffix.lower()
            media_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
                         ".gif": "image/gif", ".webp": "image/webp"}
            media_type = media_map.get(suffix, "image/jpeg")

            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=1024,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
                        {"type": "text", "text": question},
                    ],
                }],
            )
            return response.content[0].text if response.content else "No response from Claude."

        elif provider == "openai":
            api_key = config.get("openai_key", "")
            if not api_key:
                return "Error: OpenAI API key not configured."
            import openai

            client = openai.OpenAI(api_key=api_key)
            b64 = base64.standard_b64encode(image_bytes).decode("utf-8")
            suffix = path.suffix.lower()
            media_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
                         ".gif": "image/gif", ".webp": "image/webp"}
            media_type = media_map.get(suffix, "image/jpeg")

            response = client.chat.completions.create(
                model="gpt-5.2",
                max_tokens=1024,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{b64}"}},
                        {"type": "text", "text": question},
                    ],
                }],
            )
            return response.choices[0].message.content or "No response from OpenAI."

        else:
            return f"Error: unsupported provider '{provider}' for image analysis."

    except Exception as e:
        return f"Error: analyze_image failed: {type(e).__name__}: {str(e)[:200]}"


# ============================================================
# Execution
# ============================================================


def check_spending(days: int = 30) -> str:
    """Get spending summary from connected bank."""
    try:
        from plaid_integration import spending_summary, is_bank_connected
        if not is_bank_connected():
            return "No bank account connected yet. Ask the user to connect their bank through the Kiyomi app settings."
        cfg = load_config()
        plaid_cfg = cfg.get("plaid", {})
        client_id = plaid_cfg.get("client_id", "")
        secret = plaid_cfg.get("secret", "")
        env = plaid_cfg.get("env", "sandbox")
        if not client_id or not secret:
            return "Plaid is not configured yet. The user needs to add their Plaid API keys in settings."
        return spending_summary(client_id, secret, env, days)
    except ImportError:
        return "Plaid integration not available."
    except Exception as e:
        return f"Error checking spending: {e}"


def check_balances() -> str:
    """Get current bank balances."""
    try:
        from plaid_integration import balance_summary, is_bank_connected
        if not is_bank_connected():
            return "No bank account connected yet. Ask the user to connect their bank through the Kiyomi app settings."
        cfg = load_config()
        plaid_cfg = cfg.get("plaid", {})
        client_id = plaid_cfg.get("client_id", "")
        secret = plaid_cfg.get("secret", "")
        env = plaid_cfg.get("env", "sandbox")
        if not client_id or not secret:
            return "Plaid is not configured yet. The user needs to add their Plaid API keys in settings."
        return balance_summary(client_id, secret, env)
    except ImportError:
        return "Plaid integration not available."
    except Exception as e:
        return f"Error checking balances: {e}"


def check_category_spending(category: str, days: int = 30) -> str:
    """Get spending for a specific category."""
    try:
        from plaid_integration import category_spending, is_bank_connected
        if not is_bank_connected():
            return "No bank account connected yet. Ask the user to connect their bank through the Kiyomi app settings."
        cfg = load_config()
        plaid_cfg = cfg.get("plaid", {})
        client_id = plaid_cfg.get("client_id", "")
        secret = plaid_cfg.get("secret", "")
        env = plaid_cfg.get("env", "sandbox")
        if not client_id or not secret:
            return "Plaid is not configured yet. The user needs to add their Plaid API keys in settings."
        return category_spending(client_id, secret, env, category, days)
    except ImportError:
        return "Plaid integration not available."
    except Exception as e:
        return f"Error checking category spending: {e}"


# ============================================================
# Calendar tools
# ============================================================


def check_calendar(days: int = 0) -> str:
    """Get calendar events — today or upcoming N days."""
    try:
        from calendar_integration import (
            get_todays_events,
            get_upcoming_events,
            is_calendar_configured,
        )

        if not is_calendar_configured():
            return (
                "Google Calendar not connected yet. The user needs to set up "
                "Google Calendar OAuth — run setup_calendar() or add credentials "
                "to ~/.kiyomi/google_credentials.json"
            )

        if days <= 1:
            return get_todays_events()
        else:
            return get_upcoming_events(days)
    except ImportError:
        return "Calendar integration not available. Install: pip install google-auth google-auth-oauthlib google-api-python-client"
    except Exception as e:
        return f"Error checking calendar: {type(e).__name__}: {str(e)[:200]}"


def create_calendar_event(
    title: str, start: str, end: str, description: str = "", location: str = ""
) -> str:
    """Create a new calendar event."""
    try:
        from calendar_integration import create_event, is_calendar_configured

        if not is_calendar_configured():
            return (
                "Google Calendar not connected yet. The user needs to set up "
                "Google Calendar OAuth first."
            )

        return create_event(title, start, end, description, location)
    except ImportError:
        return "Calendar integration not available. Install: pip install google-auth google-auth-oauthlib google-api-python-client"
    except Exception as e:
        return f"Error creating event: {type(e).__name__}: {str(e)[:200]}"


def find_free_time_tool(date: str = "") -> str:
    """Find free time slots on a given day."""
    try:
        from calendar_integration import find_free_time, is_calendar_configured

        if not is_calendar_configured():
            return (
                "Google Calendar not connected yet. The user needs to set up "
                "Google Calendar OAuth first."
            )

        return find_free_time(date)
    except ImportError:
        return "Calendar integration not available. Install: pip install google-auth google-auth-oauthlib google-api-python-client"
    except Exception as e:
        return f"Error finding free time: {type(e).__name__}: {str(e)[:200]}"


_TOOL_FUNCS = {
    "web_search": web_search,
    "read_url": read_url,
    "read_file": read_file,
    "run_code": run_code,
    "remember": remember,
    "create_file": create_file,
    "send_email": send_email,
    "analyze_image": analyze_image,
    "check_spending": check_spending,
    "check_balances": check_balances,
    "check_category_spending": check_category_spending,
    "check_calendar": check_calendar,
    "create_event": create_calendar_event,
    "find_free_time": find_free_time_tool,
}


def execute_tool(name: str, args: dict[str, Any] | None) -> str:
    args = args or {}
    tool = _TOOL_FUNCS.get(name)
    if not tool:
        return f"Error: unknown tool '{name}'."

    try:
        if name == "web_search":
            return web_search(str(args.get("query", "")))
        if name == "read_url":
            return read_url(str(args.get("url", "")))
        if name == "read_file":
            return read_file(str(args.get("path", "")))
        if name == "run_code":
            return run_code(str(args.get("code", "")))
        if name == "remember":
            return remember(str(args.get("fact", "")), str(args.get("category", "")))
        if name == "create_file":
            return create_file(
                str(args.get("filename", "")),
                str(args.get("content", "")),
                str(args.get("title", "")),
                str(args.get("format", "general")),
            )
        if name == "send_email":
            return send_email(
                str(args.get("to", "")),
                str(args.get("subject", "")),
                str(args.get("body", "")),
            )
        if name == "analyze_image":
            return analyze_image(
                str(args.get("image_path", "")),
                str(args.get("question", "")),
            )
        if name == "check_spending":
            days = int(args.get("days", 30))
            return check_spending(days)
        if name == "check_balances":
            return check_balances()
        if name == "check_category_spending":
            return check_category_spending(
                str(args.get("category", "Food and Drink")),
                int(args.get("days", 30)),
            )
        if name == "check_calendar":
            days = int(args.get("days", 0))
            return check_calendar(days)
        if name == "create_event":
            return create_calendar_event(
                str(args.get("title", "")),
                str(args.get("start", "")),
                str(args.get("end", "")),
                str(args.get("description", "")),
                str(args.get("location", "")),
            )
        if name == "find_free_time":
            return find_free_time_tool(str(args.get("date", "")))
        return "Error: tool dispatch fell through."
    except Exception as e:
        return f"Error: tool '{name}' failed: {type(e).__name__}: {str(e)[:200]}"


# Parameters that are optional (not required)
_OPTIONAL_PARAMS = {
    "create_file": {"title", "format"},
    "check_spending": {"days"},
    "check_balances": set(),  # no params needed
    "check_category_spending": {"days"},
    "check_calendar": {"days"},
    "create_event": {"description", "location"},
    "find_free_time": {"date"},
}


def get_openai_tools_schema() -> list[dict[str, Any]]:
    """Return TOOLS converted to OpenAI 'tools' schema."""
    out: list[dict[str, Any]] = []
    for t in TOOLS:
        props = {k: {"type": v["type"], "description": v.get("description", "")} for k, v in t["parameters"].items()}
        optional = _OPTIONAL_PARAMS.get(t["name"], set())
        required = [k for k in t["parameters"] if k not in optional]
        out.append(
            {
                "type": "function",
                "function": {
                    "name": t["name"],
                    "description": t["description"],
                    "parameters": {
                        "type": "object",
                        "properties": props,
                        "required": required,
                    },
                },
            }
        )
    return out


def get_anthropic_tools_schema() -> list[dict[str, Any]]:
    """Return TOOLS converted to Anthropic 'tools' schema."""
    out: list[dict[str, Any]] = []
    for t in TOOLS:
        properties = dict(t["parameters"])
        optional = _OPTIONAL_PARAMS.get(t["name"], set())
        required = [k for k in properties if k not in optional]
        out.append(
            {
                "name": t["name"],
                "description": t["description"],
                "input_schema": {
                    "type": "object",
                    "properties": properties,
                    "required": required,
                },
            }
        )
    return out


def build_tool_schemas_gemini() -> list[dict[str, Any]]:
    """Return TOOLS converted to Gemini 'tools' schema (compatibility function)."""
    # Gemini uses a similar format to OpenAI
    return get_openai_tools_schema()

